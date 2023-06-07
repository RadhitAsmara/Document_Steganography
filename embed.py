import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor


def hide_text_in_docx(file_path, secret_text):
    doc = docx.Document(file_path)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run()
    font = run.font
    font.size = Pt(2)

    font.color.rgb = RGBColor(255, 255, 255)

    run.text = secret_text

    output_path = "modified.docx"
    doc.save(output_path)
    print(f"Steganography successfull! Output saved as {output_path}")


file_path = "original.docx"
secret_text = "This is a secret message!"

hide_text_in_docx(file_path, secret_text)
